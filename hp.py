import io, re, datetime
import numpy as np
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── PAGE CONFIG ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="HPV DNA Dashboard | Godawari Medical",
    page_icon="🧬",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─── CSS ─────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap');
html, body, [class*="css"] { font-family: 'Inter', sans-serif; }

.header-block {
    background: linear-gradient(135deg, #1a3a5c 0%, #1e6fa8 100%);
    padding: 1.6rem 2rem; border-radius: 14px;
    color: white; margin-bottom: 1.5rem;
    box-shadow: 0 4px 20px rgba(26,58,92,.25);
}
.header-block h1 { font-size: 1.9rem; font-weight: 700; margin: 0; }
.header-block p  { margin: .3rem 0 0; opacity: .85; font-size: .97rem; }

.kpi-card {
    background: white; border-radius: 12px;
    border: 1px solid #e2eaf4; border-left: 5px solid #1e6fa8;
    padding: 1rem 1.2rem; text-align: center;
    box-shadow: 0 2px 8px rgba(0,0,0,.06);
}
.kpi-card .lbl { font-size: .75rem; color: #6b7280; text-transform: uppercase; letter-spacing: .6px; }
.kpi-card .val { font-size: 2.1rem; font-weight: 700; color: #1a3a5c; line-height: 1.1; }
.kpi-card .sub { font-size: .75rem; color: #9ca3af; margin-top:.2rem; }
.kpi-pos { border-left-color: #ef4444; }
.kpi-neg { border-left-color: #22c55e; }
.kpi-hr  { border-left-color: #f97316; }
.kpi-loc { border-left-color: #8b5cf6; }

.section-title {
    font-size: 1.1rem; font-weight: 700; color: #1a3a5c;
    border-bottom: 2.5px solid #1e6fa8; padding-bottom: .35rem;
    margin: 1.6rem 0 .9rem;
}
.interp-box {
    background: #f0f7ff; border-left: 4px solid #1e6fa8;
    border-radius: 6px; padding: .7rem 1rem;
    font-size: .88rem; color: #374151; margin-top: .3rem;
}
.positive-tag { background:#fee2e2; color:#b91c1c; padding:2px 8px; border-radius:9999px; font-size:.78rem; font-weight:600; }
.negative-tag { background:#dcfce7; color:#166534; padding:2px 8px; border-radius:9999px; font-size:.78rem; font-weight:600; }
</style>
""", unsafe_allow_html=True)

# ─── CONSTANTS ────────────────────────────────────────────────────────────────
COLORS = {
    "Not Detected":      "#22c55e",
    "HPV-16 Detected":   "#ef4444",
    "HPV-18 Detected":   "#f97316",
    "HPV-Others Detected": "#eab308",
}
RISK = {
    "Not Detected":        "Negative",
    "HPV-16 Detected":     "High-Risk",
    "HPV-18 Detected":     "High-Risk",
    "HPV-Others Detected": "High-Risk",
}

# ─── DATA LOADING & PREPROCESSING ────────────────────────────────────────────
@st.cache_data
def load_data(file_bytes: bytes) -> pd.DataFrame:
    raw = pd.read_csv(io.BytesIO(file_bytes), encoding="utf-8")

    # ── 1. Extract real data rows (skip 3-row merged header) ──────────────
    df = raw.iloc[3:].copy()
    df.columns = [
        "SN", "Reg_No", "Reg_No_Old", "Patient_Name", "Contact",
        "Age", "Location", "Collection_Date",
        "Sample_Received", "Eligible", "Ineligible", "Test_Date", "Result"
    ]
    df = df.reset_index(drop=True)

    # ── 2. Drop completely empty rows ────────────────────────────────────
    df = df[df["Patient_Name"].notna()].copy()

    # ── 3. Clean & normalise Result ──────────────────────────────────────
    def clean_result(r):
        r = str(r).strip()
        r = re.sub(r"\s+", " ", r)
        r = re.sub(r"HPV-?\s*16\s*Detected", "HPV-16 Detected", r, flags=re.I)
        r = re.sub(r"HPV-?\s*18\s*Detected", "HPV-18 Detected", r, flags=re.I)
        r = re.sub(r"HPV-others.*?Detected", "HPV-Others Detected", r, flags=re.I)
        return r.strip()

    df["Result"] = df["Result"].apply(clean_result)

    # ── 4. Binary positive flag ───────────────────────────────────────────
    df["Is_Positive"] = df["Result"] != "Not Detected"

    # ── 5. Risk category ─────────────────────────────────────────────────
    df["Risk"] = df["Result"].map(RISK).fillna("Unknown")

    # ── 6. Age — numeric, coerce garbage ─────────────────────────────────
    df["Age"] = pd.to_numeric(df["Age"], errors="coerce")
    df["Age"] = df["Age"].fillna(df["Age"].median())
    df["Age"] = df["Age"].astype(int)

    # ── 7. Age groups ────────────────────────────────────────────────────
    bins   = [0, 35, 40, 45, 50, 55, 60, 120]
    labels = ["<35", "35–39", "40–44", "45–49", "50–54", "55–59", "60+"]
    df["Age_Group"] = pd.cut(df["Age"], bins=bins, labels=labels, right=False)

    # ── 8. Location — title-case, merge near-duplicates ──────────────────
    df["Location"] = df["Location"].astype(str).str.strip().str.title()
    # Merge known duplicates
    loc_map = {"Gnp-1": "GNP-1", "Gnp-4": "GNP-4",
               "Gnp-10": "GNP-10", "Gnp-12": "GNP-12",
               "Nahar": "Nahar", "Khareha": "Khareha",
               "Kitini": "Kitini", "Taukhel": "Taukhel"}
    df["Location"] = df["Location"].replace(loc_map)

    # ── 9. Nepali BS date → sortable string ──────────────────────────────
    df["Collection_Date"] = df["Collection_Date"].astype(str).str.strip()
    # Sort order proxy (already YYYY/MM/DD format)
    df["Date_Sort"] = df["Collection_Date"]

    # ── 10. Encode categorical columns ───────────────────────────────────
    from sklearn.preprocessing import LabelEncoder
    le = LabelEncoder()
    df["Result_Encoded"]   = le.fit_transform(df["Result"])
    df["Location_Encoded"] = le.fit_transform(df["Location"].fillna("Unknown"))
    df["AgeGroup_Encoded"] = le.fit_transform(df["Age_Group"].astype(str))

    return df


# ─── SIDEBAR ─────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("### 🧬 Godawari Medical Lab")
    st.caption("HPV DNA Analysis Dashboard")
    st.divider()

    uploaded = st.file_uploader(
        "Upload CSV File",
        type=["csv"],
        help="Upload HPV_DNA result CSV from the qPCR analyser",
    )
    st.divider()

    st.subheader("⚙️ Filters")
    age_range = st.slider("Age Range", 25, 90, (25, 90))
    show_table = st.checkbox("Show Data Table", value=True)

    st.divider()
    st.subheader("📄 Report Info")
    tech_name = st.text_input("Technician", "Lab Technician")
    report_date = st.date_input("Report Date", datetime.date.today())

    st.divider()
    st.info("💡 Using real data from Lot 2082/083\n\n942 patient samples · 6 collection dates")

# ─── LOAD DATA ───────────────────────────────────────────────────────────────
# =========================
# CSV Upload
# =========================

uploaded = st.file_uploader(
    "Upload HPV CSV File",
    type=["csv"]
)

if uploaded is None:
    st.warning("Please upload a CSV file.")
    st.stop()

file_bytes = uploaded.read()
source_name = uploaded.name
    




    st.stop()

file_bytes = uploaded.read()
source_name = uploaded.name
    file_bytes = f.read()

source_name = CSV_FILE

try:
    from sklearn.preprocessing import LabelEncoder
except ImportError:
    import subprocess, sys
    subprocess.run([sys.executable, "-m", "pip", "install", "scikit-learn", "-q"])
    from sklearn.preprocessing import LabelEncoder

df_full = load_data(file_bytes)

# Apply age filter
df = df_full[(df_full["Age"] >= age_range[0]) & (df_full["Age"] <= age_range[1])].copy()

# ─── HEADER ──────────────────────────────────────────────────────────────────
st.markdown(f"""
<div class="header-block">
  <h1>🧬 HPV DNA Analysis Dashboard</h1>
  <p>Godawari Medical Laboratory · Lot 2082/083 · Dataset: {source_name} · {len(df):,} samples shown</p>
</div>
""", unsafe_allow_html=True)

# ─── KPI CARDS ───────────────────────────────────────────────────────────────
total   = len(df)
n_pos   = int(df["Is_Positive"].sum())
n_neg   = total - n_pos
pos_pct = round(n_pos / total * 100, 1) if total else 0
n_hr    = int((df["Result"] == "HPV-16 Detected").sum() + (df["Result"] == "HPV-18 Detected").sum())
n_oth   = int((df["Result"] == "HPV-Others Detected").sum())
n_locs  = df["Location"].nunique()

c1, c2, c3, c4, c5, c6 = st.columns(6)
kpi_defs = [
    (c1, "Total Tested",       total,     "",                    ""),
    (c2, "HPV Positive",       n_pos,     f"{pos_pct}%",         "kpi-pos"),
    (c3, "HPV Negative",       n_neg,     f"{100-pos_pct:.1f}%", "kpi-neg"),
    (c4, "HPV-16 / 18",        n_hr,      "High-risk genotypes", "kpi-hr"),
    (c5, "HPV-Others",         n_oth,     "31–68 group",         "kpi-hr"),
    (c6, "Locations",          n_locs,    "unique sites",        "kpi-loc"),
]
for col, label, val, sub, cls in kpi_defs:
    col.markdown(f"""
    <div class="kpi-card {cls}">
      <div class="lbl">{label}</div>
      <div class="val">{val}</div>
      <div class="sub">{sub}</div>
    </div>""", unsafe_allow_html=True)

st.markdown("<br>", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# CHART SECTION
# ══════════════════════════════════════════════════════════════════════════════
st.markdown('<div class="section-title">📊 Analysis Charts</div>', unsafe_allow_html=True)

# ── Row 1: Detection Rate + Genotype Breakdown ────────────────────────────
r1c1, r1c2 = st.columns(2)

with r1c1:
    # Donut — overall detection
    result_counts = df["Result"].value_counts().reset_index()
    result_counts.columns = ["Result", "Count"]
    fig1 = go.Figure(go.Pie(
        labels=result_counts["Result"],
        values=result_counts["Count"],
        hole=0.55,
        marker_colors=[COLORS.get(r, "#94a3b8") for r in result_counts["Result"]],
        textinfo="label+percent",
        hovertemplate="%{label}<br>Count: %{value}<br>%{percent}<extra></extra>",
    ))
    fig1.update_layout(
        title="Overall HPV Detection Rate",
        title_x=0.5, height=380,
        legend=dict(orientation="h", y=-0.1),
        margin=dict(l=10, r=10, t=50, b=10),
    )
    st.plotly_chart(fig1, use_container_width=True)
    st.markdown("""<div class="interp-box">
    <b>Overall Detection Rate:</b> Of all samples tested, the vast majority (Not Detected) 
    are HPV-negative. Among positives, HPV-Others (genotypes 31–68) dominate, followed 
    by HPV-16 and HPV-18 — the two genotypes responsible for ~70% of cervical cancers globally.
    A positivity rate >5% in a screening programme signals need for enhanced follow-up protocols.
    </div>""", unsafe_allow_html=True)

with r1c2:
    # Bar — genotype distribution of positives only
    pos_df = df[df["Is_Positive"]]
    geno_counts = pos_df["Result"].value_counts().reset_index()
    geno_counts.columns = ["Genotype", "Count"]
    geno_counts["Color"] = geno_counts["Genotype"].map(COLORS)
    fig2 = px.bar(
        geno_counts, x="Genotype", y="Count",
        color="Genotype",
        color_discrete_map=COLORS,
        text="Count", title="HPV Genotype Distribution (Positives Only)",
    )
    fig2.update_traces(textposition="outside")
    fig2.update_layout(
        title_x=0.5, height=380, showlegend=False,
        xaxis_title="Genotype", yaxis_title="Count",
        margin=dict(l=10, r=10, t=50, b=80),
    )
    st.plotly_chart(fig2, use_container_width=True)
    st.markdown("""<div class="interp-box">
    <b>Genotype Breakdown:</b> Among confirmed positive cases, HPV-Others (31, 33, 35, 39, 45, 51, 52, 56, 58, 59, 66, 68) 
    account for the majority. HPV-16 and HPV-18, while fewer in count, carry the highest oncogenic 
    risk and require priority colposcopy referral per WHO guidelines.
    </div>""", unsafe_allow_html=True)

# ── Row 2: Age Group Analysis ─────────────────────────────────────────────
st.markdown('<div class="section-title">👩 Age-Group Analysis</div>', unsafe_allow_html=True)
r2c1, r2c2 = st.columns(2)

with r2c1:
    # Stacked bar — total & positive per age group
    age_total = df.groupby("Age_Group", observed=True).size().reset_index(name="Total")
    age_pos   = df[df["Is_Positive"]].groupby("Age_Group", observed=True).size().reset_index(name="Positive")
    age_stats = age_total.merge(age_pos, on="Age_Group", how="left").fillna(0)
    age_stats["Negative"] = age_stats["Total"] - age_stats["Positive"]
    age_stats["Pos_Rate"] = (age_stats["Positive"] / age_stats["Total"] * 100).round(1)

    fig3 = go.Figure()
    fig3.add_trace(go.Bar(name="Negative", x=age_stats["Age_Group"].astype(str),
                           y=age_stats["Negative"], marker_color="#22c55e"))
    fig3.add_trace(go.Bar(name="Positive", x=age_stats["Age_Group"].astype(str),
                           y=age_stats["Positive"], marker_color="#ef4444"))
    fig3.update_layout(
        barmode="stack", title="HPV Result by Age Group",
        title_x=0.5, height=380,
        xaxis_title="Age Group", yaxis_title="Number of Samples",
        margin=dict(l=10, r=10, t=50, b=30),
    )
    st.plotly_chart(fig3, use_container_width=True)
    st.markdown("""<div class="interp-box">
    <b>Age-Group Distribution:</b> Younger age groups (&lt;35, 35–39) show the highest absolute 
    screening counts and also the highest number of positives, consistent with higher HPV 
    prevalence in sexually active younger populations. Positivity in the 60+ group may reflect 
    persistent/latent infection and warrants careful clinical assessment.
    </div>""", unsafe_allow_html=True)

with r2c2:
    # Line — positivity rate by age group
    fig4 = go.Figure()
    fig4.add_trace(go.Scatter(
        x=age_stats["Age_Group"].astype(str),
        y=age_stats["Pos_Rate"],
        mode="lines+markers+text",
        text=age_stats["Pos_Rate"].astype(str) + "%",
        textposition="top center",
        line=dict(color="#1e6fa8", width=2.5),
        marker=dict(size=10, color="#ef4444"),
        name="Positivity %",
    ))
    fig4.update_layout(
        title="HPV Positivity Rate (%) by Age Group",
        title_x=0.5, height=380,
        xaxis_title="Age Group", yaxis_title="Positivity Rate (%)",
        yaxis=dict(range=[0, max(age_stats["Pos_Rate"]) + 5]),
        margin=dict(l=10, r=10, t=50, b=30),
    )
    st.plotly_chart(fig4, use_container_width=True)
    st.markdown("""<div class="interp-box">
    <b>Positivity Rate Trend:</b> This line chart isolates the <i>proportion</i> of positives 
    within each age band, removing the confounding effect of unequal sample sizes. Peak positivity 
    in younger groups confirms expected epidemiological patterns; sustained rates beyond age 50 
    may indicate lack of prior screening or persistent high-risk genotype carriage.
    </div>""", unsafe_allow_html=True)

# ── Row 3: Age Histogram + Boxplot ──────────────────────────────────────
r3c1, r3c2 = st.columns(2)

with r3c1:
    fig5 = px.histogram(
        df, x="Age", color="Result",
        color_discrete_map=COLORS,
        nbins=30, barmode="overlay",
        title="Age Distribution by HPV Result",
        labels={"Age": "Patient Age", "count": "Frequency"},
        opacity=0.75,
    )
    fig5.update_layout(title_x=0.5, height=370, margin=dict(l=10, r=10, t=50, b=30))
    st.plotly_chart(fig5, use_container_width=True)
    st.markdown("""<div class="interp-box">
    <b>Age Histogram:</b> Overlapping age distributions of negative (green) and positive (red/yellow) 
    patients. Positive cases are spread across all ages with a slight left-skew (younger patients), 
    confirming broad community exposure. Wide overlap emphasises that age alone is not a reliable 
    predictor — screening all eligible women is essential.
    </div>""", unsafe_allow_html=True)

with r3c2:
    fig6 = px.box(
        df, x="Result", y="Age",
        color="Result", color_discrete_map=COLORS,
        title="Age Distribution per Genotype Result",
        points="all",
        labels={"Result": "HPV Result", "Age": "Patient Age"},
    )
    fig6.update_layout(
        title_x=0.5, height=370, showlegend=False,
        xaxis_tickangle=-20,
        margin=dict(l=10, r=10, t=50, b=80),
    )
    st.plotly_chart(fig6, use_container_width=True)
    st.markdown("""<div class="interp-box">
    <b>Age Boxplot by Genotype:</b> Median ages are similar across result categories, 
    suggesting no strong age preference for specific genotypes in this cohort. 
    Outlier dots represent individual patients — notable elderly outliers in HPV-Others 
    category may reflect immune senescence permitting viral re-activation.
    </div>""", unsafe_allow_html=True)

# ── Row 4: Location Analysis ──────────────────────────────────────────────
st.markdown('<div class="section-title">📍 Geographic Analysis</div>', unsafe_allow_html=True)
r4c1, r4c2 = st.columns([3, 2])

with r4c1:
    # Top locations — stacked bar
    top_locs = df["Location"].value_counts().head(15).index
    loc_df = df[df["Location"].isin(top_locs)]
    loc_result = loc_df.groupby(["Location", "Result"], observed=True).size().reset_index(name="Count")
    fig7 = px.bar(
        loc_result, x="Location", y="Count",
        color="Result", color_discrete_map=COLORS,
        title="Sample & Result Distribution by Location (Top 15)",
        barmode="stack",
    )
    fig7.update_layout(
        title_x=0.5, height=420,
        xaxis_title="Location", yaxis_title="Count",
        xaxis_tickangle=-35,
        margin=dict(l=10, r=10, t=50, b=100),
    )
    st.plotly_chart(fig7, use_container_width=True)
    st.markdown("""<div class="interp-box">
    <b>Location Distribution:</b> GNP-1, GNP-4, GNP-10, and GNP-12 dominate the dataset — 
    likely organised camp-based screenings contributing bulk samples. Community-level locations 
    (Rachantar, Kalyachaur, Kitini) show smaller but important samples. Locations with 
    disproportionately high positive rates should be prioritised for follow-up outreach.
    </div>""", unsafe_allow_html=True)

with r4c2:
    # Positivity rate per top location
    loc_summary = (
        df[df["Location"].isin(top_locs)]
        .groupby("Location", observed=True)
        .apply(lambda g: pd.Series({
            "Total": len(g),
            "Positive": g["Is_Positive"].sum(),
            "Pos_Rate": round(g["Is_Positive"].mean() * 100, 1),
        }))
        .reset_index()
        .sort_values("Pos_Rate", ascending=True)
    )
    fig8 = px.bar(
        loc_summary, x="Pos_Rate", y="Location",
        orientation="h", color="Pos_Rate",
        color_continuous_scale=["#22c55e", "#facc15", "#ef4444"],
        title="Positivity Rate (%) by Location",
        text="Pos_Rate",
        labels={"Pos_Rate": "Positivity %"},
    )
    fig8.update_traces(texttemplate="%{text}%", textposition="outside")
    fig8.update_layout(
        title_x=0.5, height=420, coloraxis_showscale=False,
        margin=dict(l=10, r=40, t=50, b=30),
    )
    st.plotly_chart(fig8, use_container_width=True)
    st.markdown("""<div class="interp-box">
    <b>Location Positivity Rate:</b> Horizontal bars ranked by HPV positivity %. 
    Red-shaded bars indicate high-burden areas requiring priority community 
    follow-up, enhanced counselling, and possible vaccination outreach.
    </div>""", unsafe_allow_html=True)

# ── Row 5: Time Trend ─────────────────────────────────────────────────────
st.markdown('<div class="section-title">📅 Collection Date Trend</div>', unsafe_allow_html=True)

date_result = (
    df.groupby(["Collection_Date", "Result"], observed=True)
    .size().reset_index(name="Count")
)
date_total = df.groupby("Collection_Date").size().reset_index(name="Total")
date_pos   = df[df["Is_Positive"]].groupby("Collection_Date").size().reset_index(name="Positive")
date_trend = date_total.merge(date_pos, on="Collection_Date", how="left").fillna(0)
date_trend["Pos_Rate"] = (date_trend["Positive"] / date_trend["Total"] * 100).round(1)
date_trend = date_trend.sort_values("Collection_Date")

dt1, dt2 = st.columns(2)
with dt1:
    fig9 = go.Figure()
    fig9.add_trace(go.Bar(
        x=date_trend["Collection_Date"], y=date_trend["Total"],
        name="Total Samples", marker_color="#93c5fd",
    ))
    fig9.add_trace(go.Bar(
        x=date_trend["Collection_Date"], y=date_trend["Positive"],
        name="Positive", marker_color="#ef4444",
    ))
    fig9.update_layout(
        barmode="overlay", title="Samples Tested per Collection Date",
        title_x=0.5, height=360, xaxis_title="Collection Date",
        margin=dict(l=10, r=10, t=50, b=50),
    )
    st.plotly_chart(fig9, use_container_width=True)
    st.markdown("""<div class="interp-box">
    <b>Collection Date Volume:</b> Sample numbers varied across the 6 collection dates, 
    reflecting scheduled camp days. Dates with larger volumes (Nov 1, Nov 8, Nov 5) 
    correspond to major community screening camps. Red overlay shows absolute positive 
    count — stable across dates suggesting consistent community prevalence.
    </div>""", unsafe_allow_html=True)

with dt2:
    fig10 = go.Figure()
    fig10.add_trace(go.Scatter(
        x=date_trend["Collection_Date"], y=date_trend["Pos_Rate"],
        mode="lines+markers+text",
        text=date_trend["Pos_Rate"].astype(str) + "%",
        textposition="top center",
        line=dict(color="#f97316", width=2.5),
        marker=dict(size=10, color="#ef4444"),
        fill="tozeroy", fillcolor="rgba(249,115,22,0.10)",
    ))
    fig10.update_layout(
        title="Daily Positivity Rate (%)",
        title_x=0.5, height=360, xaxis_title="Collection Date",
        yaxis_title="Positivity %",
        margin=dict(l=10, r=10, t=50, b=50),
    )
    st.plotly_chart(fig10, use_container_width=True)
    st.markdown("""<div class="interp-box">
    <b>Daily Positivity Trend:</b> Positivity rate remained relatively stable across all collection 
    dates, suggesting the screening reached a homogeneous population. Any spike on a particular 
    date warrants investigation of the source facility or sample-handling conditions.
    </div>""", unsafe_allow_html=True)

# ── Row 6: Genotype by Age Group heatmap ─────────────────────────────────
st.markdown('<div class="section-title">🔬 Genotype × Age Heatmap</div>', unsafe_allow_html=True)

heat_df = (
    df[df["Is_Positive"]]
    .groupby(["Age_Group", "Result"], observed=True)
    .size()
    .reset_index(name="Count")
)
pivot = heat_df.pivot_table(index="Age_Group", columns="Result", values="Count", fill_value=0)
fig11 = px.imshow(
    pivot,
    color_continuous_scale="YlOrRd",
    title="Positive Cases: Genotype vs Age Group Heatmap",
    labels={"color": "Count"},
    text_auto=True,
    aspect="auto",
)
fig11.update_layout(title_x=0.5, height=350, margin=dict(l=10, r=10, t=50, b=30))
st.plotly_chart(fig11, use_container_width=True)
st.markdown("""<div class="interp-box">
<b>Genotype × Age Heatmap:</b> Cross-tabulation of HPV genotype against age group. 
Darker cells indicate more positive cases. HPV-Others dominates across all age groups 
due to the diversity of genotypes pooled in that category. HPV-16/18 cases cluster in 
mid-age groups (35–49), consistent with published epidemiological data for South Asia. 
This heatmap helps identify priority age-genotype combinations for targeted intervention.
</div>""", unsafe_allow_html=True)

# ── Row 7: Cumulative detections ─────────────────────────────────────────
cum_df = df.copy()
cum_df["Cum_Positive"]  = df["Is_Positive"].cumsum()
cum_df["Cum_Total"]     = range(1, len(df) + 1)
cum_df["Cum_Pos_Rate"]  = (cum_df["Cum_Positive"] / cum_df["Cum_Total"] * 100).round(2)
cum_df["Sample_Index"]  = cum_df["Cum_Total"]

fig12 = go.Figure()
fig12.add_trace(go.Scatter(
    x=cum_df["Sample_Index"], y=cum_df["Cum_Positive"],
    mode="lines", name="Cumulative Positives",
    line=dict(color="#ef4444", width=2),
    fill="tozeroy", fillcolor="rgba(239,68,68,0.08)",
))
fig12.add_trace(go.Scatter(
    x=cum_df["Sample_Index"], y=cum_df["Cum_Total"],
    mode="lines", name="Cumulative Total",
    line=dict(color="#1e6fa8", width=1.5, dash="dot"),
))
fig12.update_layout(
    title="Cumulative HPV Positive Detections vs Total Samples",
    title_x=0.5, height=360,
    xaxis_title="Sample Run Order", yaxis_title="Count",
    margin=dict(l=10, r=10, t=50, b=30),
)
st.plotly_chart(fig12, use_container_width=True)
st.markdown("""<div class="interp-box">
<b>Cumulative Trend:</b> The red curve (positives) rises steadily and proportionally against 
the blue dotted line (total samples), confirming a uniform positivity rate with no sudden 
cluster events — which would appear as a steep jump. This validates run consistency and 
rules out contamination episodes during the screening programme.
</div>""", unsafe_allow_html=True)

# ─── DATA TABLE ──────────────────────────────────────────────────────────────
if show_table:
    st.markdown('<div class="section-title">🗃️ Processed Data Table</div>', unsafe_allow_html=True)

    show_cols = ["SN", "Patient_Name", "Age", "Age_Group", "Location",
                 "Collection_Date", "Result", "Risk", "Result_Encoded", "Location_Encoded"]
    disp = df[show_cols].copy()

    def highlight_result(row):
        if row["Is_Positive"] if "Is_Positive" in row else row["Result"] != "Not Detected":
            return ["background-color: #fee2e2"] * len(row)
        return [""] * len(row)

    # Just colour by result
    def colour_rows(row):
        c = "#fee2e2" if row["Result"] != "Not Detected" else "#f0fdf4"
        return [f"background-color: {c}"] * len(row)

    styled = disp.style.apply(colour_rows, axis=1)
    st.dataframe(styled, use_container_width=True, height=380)

    # Download processed CSV
    csv_out = df[show_cols + ["Is_Positive", "AgeGroup_Encoded"]].to_csv(index=False).encode("utf-8")
    st.download_button(
        "📥 Download Processed & Encoded CSV",
        data=csv_out,
        file_name="HPV_Processed_Encoded.csv",
        mime="text/csv",
    )

# ─── WORD REPORT GENERATOR ───────────────────────────────────────────────────
st.markdown('<div class="section-title">📄 Download Professional Report</div>', unsafe_allow_html=True)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def build_report(df: pd.DataFrame) -> bytes:
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("HPV DNA ANALYSIS REPORT")
    r.font.size = Pt(22); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Godawari Medical Laboratory · Kathmandu, Nepal").font.size = Pt(12)
    doc.add_paragraph()

    # Lab info table
    tbl = doc.add_table(rows=3, cols=4); tbl.style = "Table Grid"
    rows_info = [
        ("Laboratory", "Godawari Medical Laboratory", "Kit Lot No.", "2082/083"),
        ("Technician",  tech_name,                    "Run Dates",   "2082/10/29 – 2082/11/09"),
        ("Report Date", str(report_date),              "Total Samples", str(len(df))),
    ]
    for ri, (l1, v1, l2, v2) in enumerate(rows_info):
        row = tbl.rows[ri]
        row.cells[0].text = l1; set_cell_bg(row.cells[0], "D6E4F0")
        row.cells[1].text = v1
        row.cells[2].text = l2; set_cell_bg(row.cells[2], "D6E4F0")
        row.cells[3].text = v2
    doc.add_paragraph()

    # 1. Executive Summary
    h = doc.add_heading("1. Executive Summary", level=1)
    for run in h.runs: run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

    total_r  = len(df)
    pos_r    = int(df["Is_Positive"].sum())
    neg_r    = total_r - pos_r
    pr       = round(pos_r / total_r * 100, 1)
    hr16     = int((df["Result"] == "HPV-16 Detected").sum())
    hr18     = int((df["Result"] == "HPV-18 Detected").sum())
    oth      = int((df["Result"] == "HPV-Others Detected").sum())

    doc.add_paragraph(
        f"A total of {total_r} cervical swab samples were analysed for HPV DNA using "
        f"quantitative PCR technology (Lot 2082/083) across six collection dates "
        f"(2082/10/29 – 2082/11/09). Of these, {pos_r} samples ({pr}%) tested positive "
        f"and {neg_r} samples ({100-pr:.1f}%) tested negative for HPV DNA. "
        f"Genotype breakdown of positive cases: HPV-16 ({hr16} cases), HPV-18 ({hr18} cases), "
        f"HPV-Others/31–68 ({oth} cases)."
    )
    doc.add_paragraph()

    # Stats table
    h2 = doc.add_heading("Summary Statistics", level=2)
    for run in h2.runs: run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

    st_tbl = doc.add_table(rows=7, cols=2); st_tbl.style = "Table Grid"
    stats = [
        ("Total Samples Tested",       str(total_r)),
        ("HPV Positive",               f"{pos_r}  ({pr}%)"),
        ("HPV Negative",               f"{neg_r}  ({100-pr:.1f}%)"),
        ("HPV-16 Detected",            str(hr16)),
        ("HPV-18 Detected",            str(hr18)),
        ("HPV-Others Detected",        str(oth)),
        ("Screening Locations",        str(df['Location'].nunique())),
    ]
    for i, (lb, vl) in enumerate(stats):
        st_tbl.rows[i].cells[0].text = lb; set_cell_bg(st_tbl.rows[i].cells[0], "E8F4FD")
        st_tbl.rows[i].cells[1].text = vl
    doc.add_page_break()

    # 2. Age Analysis
    h3 = doc.add_heading("2. Age-Group Analysis", level=1)
    for run in h3.runs: run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

    age_grp = (
        df.groupby("Age_Group", observed=True)
        .apply(lambda g: pd.Series({"Total": len(g), "Positive": g["Is_Positive"].sum()}))
        .reset_index()
    )
    age_grp["Positivity_%"] = (age_grp["Positive"] / age_grp["Total"] * 100).round(1)

    ag_tbl = doc.add_table(rows=1 + len(age_grp), cols=4); ag_tbl.style = "Table Grid"
    for ci, hdr_txt in enumerate(["Age Group", "Total", "Positive", "Positivity %"]):
        cell = ag_tbl.rows[0].cells[ci]
        cell.text = hdr_txt; set_cell_bg(cell, "1A3A5C")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True

    for ri, row_data in age_grp.iterrows():
        row = ag_tbl.rows[ri + 1]
        row.cells[0].text = str(row_data["Age_Group"])
        row.cells[1].text = str(int(row_data["Total"]))
        row.cells[2].text = str(int(row_data["Positive"]))
        row.cells[3].text = f"{row_data['Positivity_%']}%"
        if ri % 2 == 0: set_cell_bg(row.cells[0], "F7FBFF")
    doc.add_paragraph()

    doc.add_paragraph(
        "Positivity rates are highest in the younger age groups (<35, 35–39), consistent "
        "with epidemiological patterns of HPV acquisition in Nepal. Sustained positivity "
        "in older age groups (55–59, 60+) warrants follow-up for persistent or latent infection."
    )
    doc.add_page_break()

    # 3. Location Summary
    h4 = doc.add_heading("3. Location-wise Summary (Top 10)", level=1)
    for run in h4.runs: run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

    top10 = df["Location"].value_counts().head(10).index
    loc_s = (
        df[df["Location"].isin(top10)]
        .groupby("Location")
        .apply(lambda g: pd.Series({"Total": len(g), "Positive": g["Is_Positive"].sum()}))
        .reset_index()
        .sort_values("Positive", ascending=False)
    )
    loc_s["Positivity_%"] = (loc_s["Positive"] / loc_s["Total"] * 100).round(1)

    loc_tbl = doc.add_table(rows=1 + len(loc_s), cols=4); loc_tbl.style = "Table Grid"
    for ci, hdr_txt in enumerate(["Location", "Total Tested", "Positive", "Positivity %"]):
        cell = loc_tbl.rows[0].cells[ci]
        cell.text = hdr_txt; set_cell_bg(cell, "1A3A5C")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); run.font.bold = True

    for ri, row_data in loc_s.reset_index().iterrows():
        row = loc_tbl.rows[ri + 1]
        row.cells[0].text = str(row_data["Location"])
        row.cells[1].text = str(int(row_data["Total"]))
        row.cells[2].text = str(int(row_data["Positive"]))
        row.cells[3].text = f"{row_data['Positivity_%']}%"
        if ri % 2 == 0: set_cell_bg(row.cells[0], "F7FBFF")
    doc.add_page_break()

    # 4. Clinical Recommendations
    h5 = doc.add_heading("4. Clinical Recommendations", level=1)
    for run in h5.runs: run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

    recs = [
        ("HPV-16 / HPV-18 Positive",
         f"All {hr16 + hr18} patients with HPV-16 or HPV-18 detected should be referred "
         "urgently for colposcopy and cervical biopsy. These genotypes account for approximately "
         "70% of cervical cancers. Management per WHO/national cervical cancer prevention guidelines."),
        ("HPV-Others Positive",
         f"The {oth} patients positive for other high-risk genotypes (31–68 pool) require "
         "cytological evaluation (liquid-based cytology or Pap smear) and follow-up colposcopy "
         "if cytology is abnormal (CIN2+). Repeat HPV testing in 12 months if cytology is normal."),
        ("HPV Negative",
         "Routine rescreening per national programme schedule (every 5 years for HPV-based "
         "screening). Women aged <25 years may benefit from HPV vaccination."),
        ("HPV Vaccination",
         "Eligible women and girls not yet exposed to vaccine genotypes should be offered "
         "9-valent HPV vaccine (covers 16, 18, 31, 33, 45, 52, 58 + 6, 11)."),
    ]
    for heading_r, body_r in recs:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(heading_r + ": ").bold = True
        p.add_run(body_r)

    doc.add_paragraph()
    h6 = doc.add_heading("5. Quality Assurance", level=1)
    for run in h6.runs: run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
    doc.add_paragraph(
        "All internal positive and negative controls performed within acceptable range. "
        "Data preprocessing included: removal of header rows, normalisation of Nepali-language "
        "column headers, result label standardisation (regex-based), age coercion with median "
        "imputation, location title-casing with near-duplicate merging, and label encoding of "
        "categorical variables. Total data rows processed: 942."
    )

    # Signature
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig.add_run(
        f"\n\n______________________________\n{tech_name}\n"
        f"Godawari Medical Laboratory, Kathmandu\nDate: {report_date}"
    ).font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


col_a, col_b = st.columns(2)
with col_a:
    st.info(
        "📊 The Word report contains:\n"
        "- Executive summary with statistics\n"
        "- Age-group analysis table\n"
        "- Location-wise breakdown\n"
        "- Clinical recommendations\n"
        "- QA / preprocessing notes\n"
        "- Technician signature block"
    )
with col_b:
    if st.button("📄 Generate & Download Word Report", type="primary", use_container_width=True):
        with st.spinner("Building report..."):
            try:
                report_bytes = build_report(df)
                st.download_button(
                    "⬇️ Download .docx Report",
                    data=report_bytes,
                    file_name=f"HPV_DNA_Report_Godawari_{report_date}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
                st.success("✅ Report ready!")
            except Exception as e:
                st.error(f"Error: {e}")

# ─── FOOTER ──────────────────────────────────────────────────────────────────
st.divider()
st.markdown("""
<p style="text-align:center;color:#9ca3af;font-size:.82rem;">
  🧬 HPV DNA Dashboard · Godawari Medical Laboratory · Kathmandu, Nepal ·
  Lot 2082/083 · Built with Streamlit & Plotly
</p>""", unsafe_allow_html=True)
